"""
Document parsing module for extracting structured information from resumes in various formats.
Supports PDF, DOCX, HTML, and TXT formats.
"""

import os
import json
import re
from typing import Dict, Any, Optional, List
from abc import ABC, abstractmethod

# PDF parser dependencies
try:
    import fitz  # PyMuPDF
except ImportError:
    pass

# DOCX parser dependencies
try:
    import docx
except ImportError:
    pass

# HTML parser dependencies
try:
    from bs4 import BeautifulSoup
except ImportError:
    pass


class BaseResumeParser(ABC):
    """Base class for resume parsers"""
    
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse the resume file and extract structured information
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing structured resume information
        """
        pass
    
    def _extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information from text"""
        contact_info = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group(0)
        
        # Extract phone
        phone_pattern = r'\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group(0)
        
        # Extract LinkedIn (simplistic approach)
        linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9_-]+'
        linkedin_match = re.search(linkedin_pattern, text)
        if linkedin_match:
            contact_info["linkedin"] = "https://" + linkedin_match.group(0)
            
        return contact_info
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text using common skill patterns"""
        # Define lists of common programming languages, tools, frameworks
        common_skills = [
            "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Ruby", "Go", "Rust", "PHP",
            "React", "Angular", "Vue", "Node.js", "Express", "Django", "Flask", "Spring", "Ruby on Rails",
            "TensorFlow", "PyTorch", "Scikit-learn", "Pandas", "NumPy", "SQL", "NoSQL", "PostgreSQL", 
            "MySQL", "MongoDB", "AWS", "Azure", "GCP", "Docker", "Kubernetes", "CI/CD", "Git", "Linux",
            "Machine Learning", "Deep Learning", "AI", "Data Science", "DevOps", "Agile", "Scrum",
            "REST API", "GraphQL", "Microservices", "Redux", "HTML", "CSS", "SASS", "LESS"
        ]
        
        found_skills = []
        for skill in common_skills:
            skill_pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(skill_pattern, text, re.IGNORECASE):
                found_skills.append(skill)
                
        return found_skills
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        
        # Common degree patterns
        degree_patterns = [
            r'\b(Bachelor|Master|PhD|Doctorate|BSc|BA|MSc|MA|MBA|Doctor|Associate)\s+(?:of|in)?\s+([A-Za-z\s]+)',
            r'\b(B\.S\.|M\.S\.|B\.A\.|M\.A\.|M\.B\.A\.|Ph\.D\.)\s+(?:of|in)?\s+([A-Za-z\s]+)'
        ]
        
        # Common university patterns
        university_pattern = r'\b(University|College|Institute|School)\s+of\s+([A-Za-z\s]+)'
        
        # Look for degrees
        for pattern in degree_patterns:
            for match in re.finditer(pattern, text):
                degree = match.group(0)
                # Look for university names near the degree
                text_chunk = text[max(0, match.start() - 100):min(len(text), match.end() + 100)]
                university_match = re.search(university_pattern, text_chunk)
                if university_match:
                    university = university_match.group(0)
                    education.append({
                        "degree": degree,
                        "institution": university,
                        "dates": self._extract_dates_near(text_chunk, match.start() - max(0, match.start() - 100))
                    })
                
        return education
    
    def _extract_dates_near(self, text: str, offset: int = 0) -> str:
        """Extract dates near a position in text"""
        # Date patterns like 2019-2023, 2019 - 2023, Jan 2019 - Dec 2023
        date_patterns = [
            r'\b(20\d{2})\s*[-–—]\s*(20\d{2}|Present|Current|Now)\b',
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+20\d{2}\s*[-–—]\s*(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+20\d{2}\b',
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+20\d{2}\s*[-–—]\s*(Present|Current|Now)\b'
        ]
        
        for pattern in date_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                return match.group(0)
                
        return ""
    
    def _extract_work_experience(self, text: str) -> List[Dict[str, str]]:
        print(f"EXTRACT_TOP: Received text (repr, first 100 chars): {repr(text[:100])}") # DBG
        experiences = []
        
        # 1. Isolate the Work Experience section
        experience_section_text = text
        # Test with an ultra-simple header pattern (REMOVING THIS)
        # simple_header_pattern = r"^\\s*Work Experience"
        # header_match_test = re.search(simple_header_pattern, text, re.MULTILINE | re.IGNORECASE)
        # print(f"EXTRACT: Ultra-simple header test ('{simple_header_pattern}'): Match found: {bool(header_match_test)}") # DBG
        # if header_match_test:
        #     print(f"EXTRACT: Ultra-simple matched: '{header_match_test.group(0)}'")

        section_patterns = [
            # Trying to be very specific for "Work Experience" and non-greedy on the comment part
            r"^\s*Work Experience(?:\s+[^\n]*?)?\n\s*(.*?)(?:\n^\s*(?:Education|Skills|Projects|Awards|Publications|References|Languages|Contact)\s*\n|\Z)",
            # Generic one as fallback
            r"^\s*(?:Experience|Professional Experience|Employment History|Work History)(?:\s+[^\n]*?)?\n\s*(.*?)(?:\n^\s*(?:Education|Skills|Projects|Awards|Publications|References|Languages|Contact)\s*\n|\Z)",
            r"^\s*(?:Work Experience|Experience|Professional Experience|Employment History|Work History)(?:\s+[^\n]*?)?\n\s*(.*?)(?=\n\n^\s*(?:Education|Skills|Projects|Awards|Publications|References|Languages|Contact)|\Z)"
        ]
        
        found_section = False
        for pattern in section_patterns:
            section_match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE | re.DOTALL)
            print(f"EXTRACT: Trying section pattern: {pattern[:50]}... Match found: {bool(section_match)}") # DBG
            if section_match:
                experience_section_text = section_match.group(1).strip()
                print(f"EXTRACT: Isolated experience_section_text (len {len(experience_section_text)}):\n{repr(experience_section_text)[:500]}...")
                found_section = True
                break
        
        if not found_section: # If no clear section header, try to process the whole text but be more conservative
            # If no clear section header is found, return empty list.
            # The previous heuristic of checking for "experience" or "employment" keywords
            # in the whole text was too broad and led to incorrect parsing of
            # texts like the one in test_problematic_experience_parsing.
            return []

        # 2. Define regex patterns for common job elements
        # More specific date pattern, allowing for "Present", "Current", "Now"
        # Temporarily removed (?!\\s*\\() from the end to debug re.error
        date_pattern_str = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|Q[1-4]\s+\d{4})\s*(?:-|–|to|—)\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|Present|Current|Now|Q[1-4]\s+\d{4})(?!\\s*\\())"
        
        # Combined pattern for "Position at Company", "Position, Company" or "Position \\n Company"
        # This is complex. Let's try to identify blocks first.

        # Split the section into lines for easier processing
        lines = experience_section_text.split('\n')
        print(f"EXTRACT: Number of lines after split: {len(lines)}")
        if len(lines) < 5 and len(lines) > 1: # Print first few lines if few
             for idx, l_print in enumerate(lines[:5]): print(f"EXTRACT: Line {idx}: {repr(l_print)}")
        elif len(lines) == 1:
            print(f"EXTRACT: Single line content: {repr(lines[0][:500])}...")
        
        current_job_lines = []
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped: # Blank line might delineate end of a job's description
                if current_job_lines:
                    print(f"EXTRACT: Calling _parse_job_block for {len(current_job_lines)} lines ending with blank line.")
                    parsed_job = self._parse_job_block(current_job_lines, date_pattern_str)
                    if parsed_job:
                        print(f"EXTRACT: Appended job: P={parsed_job.get('position')}, C={parsed_job.get('company')}")
                        experiences.append(parsed_job)
                    else:
                        print("EXTRACT: _parse_job_block returned empty from blank line case.")
                    current_job_lines = []
                continue

            # Heuristic: If a line looks like a date range AND the *next* line is capitalized (potential new job title)
            # or if the current line is short and capitalized (potential job title) and previous block was processed.
            # This is tricky. Let's focus on parsing collected blocks.
            current_job_lines.append(line_stripped)

            # Look ahead for a very strong signal of a new job entry on the *next* line
            # e.g., if next line clearly starts with a date pattern or is a short capitalized title
            # This helps delimit current job's description.
            if i + 1 < len(lines):
                next_line_stripped = lines[i+1].strip()
                # If next line looks like a date, and current block has content, parse current.
                # Or if next line looks like a title (e.g. short, capitalized) and current_job_lines has some substance (e.g. a date found within)
                date_match_next = re.search(date_pattern_str, next_line_stripped, re.IGNORECASE)
                is_next_title_like = len(next_line_stripped.split()) < 6 and next_line_stripped == next_line_stripped.upper() and next_line_stripped # All caps
                
                # If the current line itself contains a date, and a new potential title follows, it might be the end.
                date_match_current = re.search(date_pattern_str, line_stripped, re.IGNORECASE)

                if date_match_next and len(current_job_lines) > 1 : # Next line is a date, implies new job
                     # check if current_job_lines already has a date, if so, it's a self-contained job ending here
                    has_date_already = any(re.search(date_pattern_str, l, re.IGNORECASE) for l in current_job_lines[:-1]) # exclude current line if it's the one with the date
                    if has_date_already or not date_match_current : # if current line is not the date line OR prev lines had a date
                        print(f"EXTRACT: Calling _parse_job_block for {len(current_job_lines)} lines due to lookahead.")
                        parsed_job = self._parse_job_block(current_job_lines, date_pattern_str)
                        if parsed_job:
                            print(f"EXTRACT: Appended job from lookahead: P={parsed_job.get('position')}, C={parsed_job.get('company')}")
                            experiences.append(parsed_job)
                        else:
                            print("EXTRACT: _parse_job_block returned empty from lookahead case.")
                        current_job_lines = []
                        continue # current line will be processed with next block


            # If we have accumulated a lot of lines without parsing, and then hit a blank line or end.
            # This is a fallback.

        # Process any remaining lines in current_job_lines
        if current_job_lines:
            print(f"EXTRACT: Calling _parse_job_block for remaining {len(current_job_lines)} lines.")
            parsed_job = self._parse_job_block(current_job_lines, date_pattern_str)
            if parsed_job:
                print(f"EXTRACT: Appended job from remaining: P={parsed_job.get('position')}, C={parsed_job.get('company')}")
                experiences.append(parsed_job)
            else:
                print("EXTRACT: _parse_job_block returned empty from remaining lines case.")
        
        # Filter out entries that are clearly too sparse (e.g., only got a date)
        experiences = [exp for exp in experiences if exp.get("position") != "N/A" or exp.get("company") != "N/A"]
        return experiences

    def _parse_job_block(self, job_lines: List[str], date_pattern_str: str) -> Dict[str, str]:
        print(f"PARSE_BLOCK: Received job_lines: {job_lines}")
        if not job_lines:
            print("PARSE_BLOCK: Empty job_lines, returning empty.")
            return {}

        position = "N/A"
        company = "N/A"
        dates = "N/A"
        description_list = []
        
        date_found = False
        date_line_idx = -1

        # First pass: Find the date line
        for i, line in enumerate(job_lines):
            date_match = re.search(date_pattern_str, line, re.IGNORECASE)
            if date_match:
                dates = f"{date_match.group(1).strip()} - {date_match.group(2).strip()}"
                date_found = True
                date_line_idx = i
                # Attempt to extract company from the same line if possible, or line before
                # e.g. "Company Name | Jan 2020 - Dec 2020"
                #      "Some Text Jan 2020 - Dec 2020 at Company Name"
                #      "Jan 2020 - Dec 2020, Company Name"
                
                # Company on the same line, before dates. E.g., "My Company | Location | Jan 2020 - Dec 2020"
                company_candidate_text = line[:date_match.start()].strip()
                # Try to split by " | " to separate company from potential location
                parts = [p.strip() for p in company_candidate_text.split('|') if p.strip()]
                if parts:
                    potential_company_name = parts[0].strip(" ,-@")
                    if potential_company_name and len(potential_company_name.split()) < 7 and not any(kw in potential_company_name.lower() for kw in ["experience", "details"]):
                         # Check if it's not just a job title from the previous line
                        if i > 0 and potential_company_name.lower() == job_lines[i-1].strip().lower():
                            pass # it's likely the position repeated
                        else:
                            if company == "N/A": # Prioritize if company not already found by other means (e.g. 'at Company')
                                company = potential_company_name
                
                # Company on the same line, after dates. E.g., "Jan 2020 - Dec 2020 at My Company"
                company_candidate_after_dates = line[date_match.end():].strip(" |,-@").strip()
                if company_candidate_after_dates.startswith("at "): # Common pattern
                    company_candidate_after_dates = company_candidate_after_dates[3:].strip()

                if company_candidate_after_dates and len(company_candidate_after_dates.split()) < 7 and company == "N/A":
                     company = company_candidate_after_dates
                break # Found primary date line

        if not date_found: # If no standard date, maybe it's a single year or unparsable.
                           # For simplicity, if no date, this block might be invalid or just description.
            # Try a simpler "at Company" or "Company Name" on first few lines if no date was found.
            # This makes it hard to distinguish from description.
            pass


        # Second pass: Determine position and company based on date line
        # Position is likely the line(s) before the date line, or the first line of the block.
        # Company might be on the date line, or line before/after position.
        
        if date_line_idx == 0: # Date is on the first line
            # Position might be N/A or needs more complex inference.
            # Company might be on this line (already handled) or also N/A.
            # Description starts from next line.
            if company == "N/A" and len(job_lines) > 1: # If company not found on date line, check next line
                # This is risky, could be start of description. Only if it's short.
                if len(job_lines[1].split()) < 5:
                    company = job_lines[1] # Assume next line is company
                    description_list = job_lines[2:]
                else:
                    description_list = job_lines[1:]
            else:
                 description_list = job_lines[1:]


        elif date_line_idx > 0: # Date is on a later line
            # Line before date_line_idx is often Company or Position
            # Line 0 is often Position
            potential_pos_company_lines = job_lines[:date_line_idx]
            description_list = job_lines[date_line_idx+1:]

            if len(potential_pos_company_lines) == 1:
                # Assume it's Position, Company might have been on date line or needs to be found
                line_content = potential_pos_company_lines[0]
                # Try "Position at Company" or "Position, Company"
                match_pac = re.match(r"(.+?)(?:\\s+at\\s+|\\s*,\\s*)(.+)", line_content, re.IGNORECASE)
                if match_pac:
                    position = match_pac.group(1).strip()
                    if company == "N/A": company = match_pac.group(2).strip() # Prioritize company from date line
                else: # Assume it's just position
                    position = line_content
                    # Company might be on the date line (already extracted potentially) or still N/A

            elif len(potential_pos_company_lines) > 1:
                # Typically: Line 0 = Position, Line 1 = Company
                position = potential_pos_company_lines[0]
                if company == "N/A": # If not found on date line
                    company = potential_pos_company_lines[1]
                # Refine: if line 1 looks like location, company might be line 0 if position contains "at"
                match_pos_at = re.match(r"(.+?)\\s+at\\s+(.+)", position, re.IGNORECASE)
                if match_pos_at:
                    position = match_pos_at.group(1).strip()
                    if company == "N/A": company = match_pos_at.group(2).strip()
        
        else: # date_line_idx == -1 (no date found using primary pattern)
              # Fallback: assume first line is pos, second is company, rest is desc.
              # This is very basic.
            if len(job_lines) > 0:
                pos_candidate = job_lines[0]
                match_pac = re.match(r"(.+?)(?:\\s+at\\s+|\\s*,\\s*)(.+)", pos_candidate, re.IGNORECASE)
                if match_pac:
                    position = match_pac.group(1).strip()
                    company = match_pac.group(2).strip()
                    if len(job_lines) > 1: description_list = job_lines[1:]
                elif len(job_lines) > 1:
                    position = pos_candidate
                    company_candidate = job_lines[1]
                    # Avoid taking long description as company
                    if len(company_candidate.split()) < 7:
                        company = company_candidate
                        if len(job_lines) > 2: description_list = job_lines[2:]
                    else: # Assume line 1 is start of description
                        description_list = job_lines[1:]
                else: # Only one line, assume it's position or company
                    position = pos_candidate # Or company = pos_candidate
                    description_list = []
            else: # No lines
                return {}


        # Clean up description
        description = "\\n".join(line.strip() for line in description_list if line.strip() and not line.strip().lower().startswith("keywords:")).strip()
        
        # Final check: if position looks like a company and company is N/A, swap
        if company == "N/A" and position != "N/A" and ("inc" in position.lower() or "llc" in position.lower() or "ltd" in position.lower() or "group" in position.lower()):
            company = position
            position = "N/A" # Or try to find a better position

        # Avoid returning if essentials are missing
        if position == "N/A" and company == "N/A" and dates == "N/A":
            print(f"PARSE_BLOCK: All essentials N/A. Pos: {position}, Comp: {company}, Dates: {dates}. Returning empty.")
            return {}
            
        print(f"PARSE_BLOCK: Returning: P={position}, C={company}, D={dates}, Desc_len={len(description)}")
        return {
            "position": position.strip(":, "),
            "company": company.strip(":, "),
            "dates": dates.strip(":, "),
            "description": description
        }


class PDFResumeParser(BaseResumeParser):
    """Parser for PDF resumes using PyMuPDF"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            full_text = ""
            doc = fitz.open(file_path)
            
            # Extract full text from all pages
            for page in doc:
                full_text += page.get_text()
                
            # Get basic contact info
            contact_info = self._extract_contact_info(full_text)
            
            # Get name from first page
            name = self._extract_name(full_text)
            
            # Extract skills
            skills = self._extract_skills(full_text)
            
            # Extract education
            education = self._extract_education(full_text)
            
            # Extract work experience
            experience = self._extract_work_experience(full_text)
            
            # Build the structured resume
            resume = {
                "name": name,
                "contact_information": contact_info,
                "skills": skills,
                "education": education,
                "work_experience": experience,
                "full_text": full_text  # Include full text for further processing if needed
            }
            
            return resume
            
        except Exception as e:
            return {"error": f"Failed to parse PDF: {str(e)}"}
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from beginning of resume"""
        # Name is often the first line or close to the top
        lines = text.strip().split('\n')
        # Use the first line that's not empty and looks like a name
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if len(line) > 0 and len(line.split()) <= 4 and line.split()[0][0].isupper():
                # Simple heuristic: first 1-4 words, first letter is uppercase
                return line
        return "Unknown"


class DOCXResumeParser(BaseResumeParser):
    """Parser for DOCX resumes using python-docx"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = docx.Document(file_path)
            full_text = ""
            
            # Extract all text
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
                
            # Get basic contact info
            contact_info = self._extract_contact_info(full_text)
            
            # Get name from the beginning
            name = self._extract_name(doc, full_text)
            
            # Extract skills
            skills = self._extract_skills(full_text)
            
            # Extract education
            education = self._extract_education(full_text)
            
            # Extract work experience
            experience = self._extract_work_experience(full_text)
            
            # Build the structured resume
            resume = {
                "name": name,
                "contact_information": contact_info,
                "skills": skills,
                "education": education,
                "work_experience": experience,
                "full_text": full_text
            }
            
            return resume
            
        except Exception as e:
            return {"error": f"Failed to parse DOCX: {str(e)}"}
    
    def _extract_name(self, doc, text: str) -> str:
        """Extract candidate name from beginning of resume"""
        # Try to get it from the first paragraph with larger font
        for paragraph in doc.paragraphs[:10]:  # Check first 10 paragraphs
            if paragraph.text.strip() and len(paragraph.text.split()) <= 4:
                # Check if this has a larger font than standard
                for run in paragraph.runs:
                    if run.font.size and run.font.size > 12:  # Font larger than 12pt
                        return paragraph.text.strip()
        
        # Fallback: use first line
        lines = text.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 0 and len(line.split()) <= 4 and line.split()[0][0].isupper():
                return line
                
        return "Unknown"


class HTMLResumeParser(BaseResumeParser):
    """Parser for HTML resumes using BeautifulSoup"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
                
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
                
            # Get text
            full_text = soup.get_text(separator='\n')
            
            # Basic info
            contact_info = self._extract_contact_info(full_text)
            
            # Try to get name from title or h1
            name = self._extract_name(soup, full_text)
            
            # Extract skills
            skills = self._extract_skills(full_text)
            
            # Extract education
            education = self._extract_education(full_text)
            
            # Extract work experience
            experience = self._extract_work_experience(full_text)
            
            # Build the structured resume
            resume = {
                "name": name,
                "contact_information": contact_info,
                "skills": skills,
                "education": education,
                "work_experience": experience,
                "full_text": full_text
            }
            
            return resume
            
        except Exception as e:
            return {"error": f"Failed to parse HTML: {str(e)}"}
    
    def _extract_name(self, soup, text: str) -> str:
        """Extract candidate name from HTML"""
        # Check for name in title
        title = soup.find('title')
        if title and title.text.strip():
            return title.text.strip()
            
        # Check h1 tags
        h1 = soup.find('h1')
        if h1 and h1.text.strip():
            return h1.text.strip()
        
        # Fallback to first lines of text
        lines = text.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 0 and len(line.split()) <= 4 and line.split()[0][0].isupper():
                return line
                
        return "Unknown"


class TXTResumeParser(BaseResumeParser):
    """Parser for plain text resumes"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                full_text = file.read()
                
            # Get basic contact info
            contact_info = self._extract_contact_info(full_text)
            
            # Get name from the beginning
            name = self._extract_name(full_text)
            
            # Extract skills
            skills = self._extract_skills(full_text)
            
            # Extract education
            education = self._extract_education(full_text)
            
            # Extract work experience
            experience = self._extract_work_experience(full_text)
            
            # Build the structured resume
            resume = {
                "name": name,
                "contact_information": contact_info,
                "skills": skills,
                "education": education,
                "work_experience": experience,
                "full_text": full_text
            }
            
            return resume
            
        except Exception as e:
            return {"error": f"Failed to parse TXT: {str(e)}"}
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from text resume"""
        lines = text.strip().split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if len(line) > 0 and len(line.split()) <= 4 and line.split()[0][0].isupper():
                return line
        return "Unknown"


def get_resume_parser(file_path: str) -> BaseResumeParser:
    """
    Factory method to get the appropriate parser based on file extension
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        An instance of the appropriate parser
    
    Raises:
        ValueError: If the file format is not supported
    """
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    
    if extension == '.pdf':
        return PDFResumeParser()
    elif extension == '.docx':
        return DOCXResumeParser()
    elif extension in ['.html', '.htm']:
        return HTMLResumeParser()
    elif extension == '.txt':
        return TXTResumeParser()
    else:
        raise ValueError(f"Unsupported file format: {extension}")


def parse_resume(file_path: str) -> Dict[str, Any]:
    """
    Parse a resume file and return structured information
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Dictionary containing structured resume information
    """
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}
        
    try:
        parser = get_resume_parser(file_path)
        return parser.parse(file_path)
    except ValueError as e:
        return {"error": str(e)}
    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}"} 